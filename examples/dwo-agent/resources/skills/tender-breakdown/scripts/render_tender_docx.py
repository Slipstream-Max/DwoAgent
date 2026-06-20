#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
把招标拆解 Markdown 文档渲染为 .docx。

用法:
    uv run --with python-docx python scripts/render_tender_docx.py docs-md -o docs-docx
    uv run --with python-docx python scripts/render_tender_docx.py 招标分析文档.md -o 招标分析文档.docx
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:
    raise SystemExit(
        "缺少 python-docx。请使用："
        "uv run --with python-docx python scripts/render_tender_docx.py <input.md|dir> -o <output>"
    ) from exc


def set_east_asian_font(run_or_font, font_name: str = "宋体") -> None:
    font = getattr(run_or_font, "font", run_or_font)
    font.name = font_name

    element = getattr(run_or_font, "element", None)
    if element is not None and hasattr(element, "get_or_add_rPr"):
        r_pr = element.get_or_add_rPr()
    else:
        r_pr = getattr(font, "_element", None)

    if r_pr is not None:
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
            r_pr.append(r_fonts)
        else:
            r_fonts.set(qn("w:eastAsia"), font_name)


def set_cell_text(cell, text: str, bold: bool = False, header: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    set_east_asian_font(run, "黑体" if header else "宋体")
    if header:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F4E79"/>')
        cell._tc.get_or_add_tcPr().append(shading)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(11)
    set_east_asian_font(normal.font, "宋体")
    normal.paragraph_format.line_spacing = 1.35


def add_heading(doc: Document, text: str, level: int) -> None:
    level = min(max(level, 1), 4)
    paragraph = doc.add_paragraph()
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    font_size = {1: 18, 2: 15, 3: 13, 4: 12}[level]
    run.font.size = Pt(font_size)
    set_east_asian_font(run, "黑体")


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    set_east_asian_font(run, "宋体")


def add_list_item(doc: Document, text: str, numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    try:
        paragraph = doc.add_paragraph(style=style)
    except KeyError:
        paragraph = doc.add_paragraph()
        text = ("1. " if numbered else "- ") + text
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    set_east_asian_font(run, "宋体")


def is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        if is_table_separator(line):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    width = max((len(row) for row in rows), default=0)
    return [row + [""] * (width - len(row)) for row in rows if width]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row in enumerate(rows):
        for col_index, text in enumerate(row):
            set_cell_text(table.rows[row_index].cells[col_index], text, bold=row_index == 0, header=row_index == 0)
    doc.add_paragraph()


def collect_table_block(lines: list[str], start: int) -> tuple[list[str], int]:
    block: list[str] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|") and lines[index].strip().endswith("|"):
        block.append(lines[index])
        index += 1
    return block, index


def render_markdown_to_docx(input_path: Path, output_path: Path) -> Path:
    doc = Document()
    configure_document(doc)
    lines = input_path.read_text(encoding="utf-8-sig").splitlines()

    index = 0
    in_code_block = False
    code_lines: list[str] = []

    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                if code_lines:
                    add_paragraph(doc, "\n".join(code_lines))
                    code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            index += 1
            continue

        if in_code_block:
            code_lines.append(raw_line)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_block, index = collect_table_block(lines, index)
            add_table(doc, parse_table(table_block))
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            add_heading(doc, heading_match.group(2).strip(), len(heading_match.group(1)))
            index += 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet_match:
            add_list_item(doc, bullet_match.group(1).strip())
            index += 1
            continue

        number_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if number_match:
            add_list_item(doc, number_match.group(1).strip(), numbered=True)
            index += 1
            continue

        add_paragraph(doc, stripped)
        index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def discover_inputs(items: list[str]) -> list[Path]:
    paths: list[Path] = []
    for item in items:
        path = Path(item).expanduser().resolve()
        if path.is_dir():
            paths.extend(sorted(path.glob("*.md")))
        elif path.is_file() and path.suffix.lower() == ".md":
            paths.append(path)
        else:
            raise RuntimeError(f"不是 Markdown 文件或目录：{path}")
    return paths


def output_path_for(input_path: Path, output_arg: str | None, multiple: bool) -> Path:
    if output_arg is None:
        return input_path.with_suffix(".docx")
    output = Path(output_arg).expanduser().resolve()
    if multiple or output.suffix.lower() != ".docx":
        return output / f"{input_path.stem}.docx"
    return output


def main() -> int:
    parser = argparse.ArgumentParser(description="把招标拆解 Markdown 文档渲染为 Word docx")
    parser.add_argument("inputs", nargs="+", help="Markdown 文件或包含 Markdown 的目录")
    parser.add_argument("-o", "--output", help="输出 docx 路径；批量时作为输出目录")
    args = parser.parse_args()

    try:
        inputs = discover_inputs(args.inputs)
        if not inputs:
            raise RuntimeError("未找到 Markdown 文件")
        multiple = len(inputs) > 1 or any(Path(item).expanduser().resolve().is_dir() for item in args.inputs)
        for input_path in inputs:
            output_path = output_path_for(input_path, args.output, multiple)
            result = render_markdown_to_docx(input_path, output_path)
            print(f"已生成：{result}")
    except Exception as exc:  # noqa: BLE001
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
